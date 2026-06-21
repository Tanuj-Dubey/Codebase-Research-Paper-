import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from xhtml2pdf import pisa

# --- Configuration ---
PROJECT_NAME = "Hybrid Slm-RUL Framework Research"
PLOTS_DIR = "plots"
HTML_SOURCE = "Final_Research_Report.html"
DOCX_OUT = "Battery_RUL_Final_Report.docx"
PDF_OUT = "Battery_RUL_Final_Report.pdf"

# --- Content Definitions (Same to Same) ---
sections = [
    {
        "title": "I. Proposed Hybrid SLM-RUL Framework",
        "text": "The framework synergizes the generalized pattern recognition of Amazon Chronos with a localized Small Language Model (SLM) trained specifically on generative residual errors. Post-hoc calibration via Ridge-Affine mapping ensures peak precision across varying battery chemistries.",
        "image": "architecture_diagram.png",
        "caption": "Fig. 1. Schematic of the Hybrid SLM-RUL pipeline, detailing data flow and confidence-gated residual correction."
    },
    {
        "title": "II. Data Manifold Validation (t-SNE)",
        "text": "Generative data augmentation was utilized to address cycle-life data scarcity. t-SNE manifold projections verify that synthetic features effectively occupy the operational feature space of true operational cells.",
        "image": "B0005_tsne.png",
        "caption": "Fig. 2. t-SNE projection of concatenated real and synthetic datasets for battery cell B0005."
    },
    {
        "title": "III. Training Convergence Dynamics (Case 1)",
        "text": "Evaluates temporal generalization on individual trajectories. The optimization objective shows stable blue (train) and green (val) trajectories, verifying efficient feature extraction.",
        "image": "loss_case1.png",
        "caption": "Fig. 3. Training and Validation trajectories under the Case 1 (Intra-cell) split."
    },
    {
        "title": "IV. Cross-Battery Stability (Case 3 - LOO)",
        "text": "The Leave-One-Out (LOO) test on Case 3 confirms the model's capacity to maintain a low error gap even when forecasting for completely unseen cell signatures.",
        "image": "loss_case3.png",
        "caption": "Fig. 4. Convergence profiles for the cross-battery Leave-One-Out experimental setup."
    },
    {
        "title": "V. Technical Specifications & Hyperparameters",
        "text": "The table below outlines the converged configuration settings for the SLM Transformer, PatchTST Residuals, and Chronos Baseline forensic settings.",
        "image": "comprehensive_hyperparameters.png",
        "caption": "Table 1. Unified hyperparameter matrix across all experimental boundary cases."
    },
    {
        "title": "VI. Integrated Remaining Useful Life Forecasts",
        "text": "Observed vs. Predicted discharge capacity profiles. The system effectively tracks non-linear degradation and stays within tight Ah tolerance boundaries even as the cell approaches End-of-Life (EOL).",
        "image": "rul_case1_B0005.png",
        "caption": "Fig. 5. Comparative RUL predictions for cell B0005. The transition line denotes the autonomous forecasting start."
    },
    {
        "title": "VII. Methodological Implementation",
        "text": "Standardized computational algorithm for the implementation of the Hybrid Slm-RUL framework.",
        "image": "algorithm_pseudocode.png",
        "caption": "Algorithm 1. Hybrid Inference Methodology for battery health prognosis."
    }
]

# --- 1. DOCX Generation ---
def generate_docx():
    print(f"Generating {DOCX_OUT}...")
    doc = Document()
    
    # Title
    title = doc.add_heading("Hybrid Generative Residual Networks for Battery RUL Forecasting", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph("Technical Research Summary | Project Results Submission")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for sec in sections:
        doc.add_heading(sec["title"], level=1)
        doc.add_paragraph(sec["text"])
        
        img_path = os.path.join(PLOTS_DIR, sec["image"])
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(5.5))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            cap = doc.add_paragraph(sec["caption"])
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.style.font.size = Pt(10)
            cap.style.font.italic = True
        
        doc.add_page_break()
        
    doc.save(DOCX_OUT)
    return True

# --- 2. PDF Generation ---
def generate_pdf():
    print(f"Generating {PDF_OUT}...")
    with open(HTML_SOURCE, "r", encoding="utf-8") as f:
        html = f.read()
    
    # xhtml2pdf works best with local relative paths or absolute file paths
    # We'll use the same relative paths used in Final_Research_Report.html
    
    with open(PDF_OUT, "wb") as f:
        pisa_status = pisa.CreatePDF(html, dest=f)
    
    return not pisa_status.err

# --- Run ---
if __name__ == "__main__":
    if generate_docx():
        print(f"Successfully created {DOCX_OUT}")
    if generate_pdf():
        print(f"Successfully created {PDF_OUT}")
    else:
        print("PDF generation encountered errors.")
